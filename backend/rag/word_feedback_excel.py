import re
from datetime import datetime, timedelta
from pathlib import Path

import pandas as pd
from docx import Document


def parse_date(date_str):
    """Parse dates from multiple formats and return datetime or None."""
    cleaned = date_str.strip().replace("-", "/").replace(".", "/")
    for fmt in ("%d/%m/%Y", "%m/%d/%Y", "%d/%m/%y", "%m/%d/%y", "%Y/%m/%d"):
        try:
            dt = datetime.strptime(cleaned, fmt)
            # Force 2-digit year values to 2000+ (e.g., 23 -> 2023).
            if dt.year < 100:
                dt = dt.replace(year=2000 + dt.year)
            return dt
        except ValueError:
            continue
    return None


def format_grade_or_year(num, level_type):
    if 11 <= (num % 100) <= 13:
        suffix = "th"
    else:
        suffix = {1: "st", 2: "nd", 3: "rd"}.get(num % 10, "th")
    return f"{num}{suffix} {level_type.title()}"


def normalize_school_name(raw_name):
    """Map noisy school names to canonical values."""
    if not raw_name:
        return "Unknown School"
    text = re.sub(r"\s+", " ", raw_name).strip()
    low = text.lower()
    mapping = [
        ("isaiambalam", "Isaiambalam"),
        ("isai", "Isai"),
        ("udavi", "Udavi"),
        ("aiat", "AIAT"),
        ("aikiyam", "Aikiyam"),
        ("future", "Future"),
        ("government school", "Government School"),
        ("govt school", "Government School"),
        ("government", "Government School"),
        ("nes", "NES"),
        ("stem intuition", "Stem Intuition"),
        ("stem", "Stem Intuition"),
        ("transition", "Transition"),
        ("vasavi", "Vasavi"),
    ]
    for key, canonical in mapping:
        if key in low:
            return canonical
    return text.title()


def format_date_for_excel(value):
    """Convert date string/timestamp to dd-mm-yyyy."""
    if pd.isna(value):
        return ""
    dt = pd.to_datetime(value, errors="coerce")
    if pd.isna(dt):
        return ""
    return dt.strftime("%d-%m-%Y")


def extract_grade_from_text(line):
    """Handle many grade/year formats used in notes."""
    pattern = re.compile(
        r"(?:\b(grade|year|std|standard|class)\b\s*[:\-]?\s*0*(\d+)(?:st|nd|rd|th)?|"
        r"\b0*(\d+)(?:st|nd|rd|th)?\s*(grade|year|std|standard|class)\b)",
        re.IGNORECASE,
    )
    match = pattern.search(line)
    if not match:
        return None

    level_type = (match.group(1) or match.group(4) or "grade").lower()
    level_type = "Year" if level_type == "year" else "Grade"
    grade_num = int(match.group(2) or match.group(3))
    return format_grade_or_year(grade_num, level_type)


def grade_sort_key(value):
    if not isinstance(value, str):
        return 999
    match = re.search(r"(\d+)", value)
    return int(match.group(1)) if match else 999


def extract_lines_from_docx(file_path):
    doc = Document(file_path)
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if text:
                        lines.append(text)
    return lines


def extract_date_range_from_filename(file_name):
    # Numeric examples: 1_06_23_31_12_23 / 01-01-2023 to 31-08-2023
    pattern = re.compile(
        r"(\d{1,2})[\/_.-](\d{1,2})[\/_.-](\d{2,4}).*?(\d{1,2})[\/_.-](\d{1,2})[\/_.-](\d{2,4})",
        re.IGNORECASE,
    )
    match = pattern.search(file_name)
    if match:
        s_day, s_month, s_year, e_day, e_month, e_year = match.groups()
        start_dt = parse_date(f"{s_day}/{s_month}/{s_year}")
        end_dt = parse_date(f"{e_day}/{e_month}/{e_year}")
        return start_dt, end_dt

    # Month-name examples: Jan_June_2025 / June_December_2025
    month_name_pattern = re.compile(
        r"(jan|january|feb|february|mar|march|apr|april|may|jun|june|jul|july|aug|august|sep|sept|september|oct|october|nov|november|dec|december)"
        r".*?"
        r"(jan|january|feb|february|mar|march|apr|april|may|jun|june|jul|july|aug|august|sep|sept|september|oct|october|nov|november|dec|december)"
        r".*?"
        r"(20\d{2})",
        re.IGNORECASE,
    )
    month_map = {
        "jan": 1,
        "january": 1,
        "feb": 2,
        "february": 2,
        "mar": 3,
        "march": 3,
        "apr": 4,
        "april": 4,
        "may": 5,
        "jun": 6,
        "june": 6,
        "jul": 7,
        "july": 7,
        "aug": 8,
        "august": 8,
        "sep": 9,
        "sept": 9,
        "september": 9,
        "oct": 10,
        "october": 10,
        "nov": 11,
        "november": 11,
        "dec": 12,
        "december": 12,
    }
    month_match = month_name_pattern.search(file_name)
    if not month_match:
        return None, None

    sm, em, year = month_match.groups()
    start_month = month_map[sm.lower()]
    end_month = month_map[em.lower()]
    y = int(year)
    start_dt = datetime(y, start_month, 1)
    # Last day of end month
    if end_month == 12:
        end_dt = datetime(y, 12, 31)
    else:
        end_dt = datetime(y, end_month + 1, 1) - timedelta(days=1)
    return start_dt, end_dt


def process_feedback_text(file_path):
    data = []
    lines = extract_lines_from_docx(file_path)

    current_start_date = None
    current_end_date = None
    current_school = "Unknown School"
    current_grade = None
    current_teachers = None
    current_notes_count = None
    current_reported_by = None
    current_feedback = []

    date_pattern = re.compile(r"Date:\s*(\d{1,2}[\/\-.]\d{1,2}[\/\-.]\d{2,4})", re.IGNORECASE)
    start_date_pattern = re.compile(r"Start\s*Date:\s*(\d{1,2}[\/\-.]\d{1,2}[\/\-.]\d{2,4})", re.IGNORECASE)
    end_date_pattern = re.compile(r"End\s*Date:\s*(\d{1,2}[\/\-.]\d{1,2}[\/\-.]\d{2,4})", re.IGNORECASE)
    school_pattern = re.compile(
        r"^\s*(?:School(?:\s*Name)?|College(?:\s*Name)?|Institution|Center|Centre|Campus|Unit|Location)\s*[:\-]\s*(.*)",
        re.IGNORECASE,
    )
    school_inline_pattern = re.compile(
        r"\b(?:isaiambalam|isai|udavi|aiat|aikiyam|future|government school|govt school|nes|stem intuition|transition|vasavi)\b",
        re.IGNORECASE,
    )
    grade_pattern = re.compile(
        r"^\s*(?:(grade|year)\s*0*(\d+)(?:st|nd|rd|th)?|0*(\d+)(?:st|nd|rd|th)?\s*(grade|year))[\s,:-]*(.*)",
        re.IGNORECASE,
    )
    notes_count_pattern = re.compile(r"(?:No\.?\s*of\s*notes|Notes?\s*Count)\s*[:\-]\s*(\d+)", re.IGNORECASE)
    reported_by_pattern = re.compile(r"(?:User\s*Name|Reported\s*By|Prepared\s*By)\s*[:\-]\s*(.+)", re.IGNORECASE)
    teacher_line_pattern = re.compile(r"(?:Teacher(?:s)?|Faculty)\s*[:\-]\s*(.+)", re.IGNORECASE)
    ignore_phrases = ["dear all", "truth and love", "sanjeev", "sub:"]

    def append_current_block():
        has_dates = bool(current_start_date or current_end_date)
        has_school = bool(current_school)
        has_grade = bool(current_grade and current_grade != "Unknown Grade/Year")
        has_feedback = bool("\n".join(current_feedback).strip())
        # Keep only meaningful rows, avoid empty/unknown rows.
        if has_dates and has_school and has_feedback:
            data.append(
                [
                    current_start_date,
                    current_end_date,
                    current_school,
                    current_grade if current_grade else "",
                    current_teachers,
                    "\n".join(current_feedback),
                ]
            )

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            continue

        lower_line = line.lower()
        if any(lower_line.startswith(phrase) for phrase in ignore_phrases) or "http" in lower_line:
            continue

        notes_match = notes_count_pattern.search(line)
        if notes_match:
            current_notes_count = int(notes_match.group(1))
            continue

        reported_by_match = reported_by_pattern.search(line)
        if reported_by_match:
            current_reported_by = reported_by_match.group(1).strip()
            continue

        teacher_line_match = teacher_line_pattern.search(line)
        if teacher_line_match and not current_teachers:
            current_teachers = teacher_line_match.group(1).strip()
            continue

        school_match = school_pattern.match(line)
        if school_match:
            append_current_block()
            current_school = normalize_school_name(school_match.group(1).strip())
            current_grade = None
            current_teachers = None
            current_feedback = []
            continue

        # fallback school detection when label is missing: only for short, title-like lines
        cleaned_inline = re.sub(r"[^a-zA-Z ]", " ", line).strip()
        if (
            school_inline_pattern.search(cleaned_inline)
            and len(cleaned_inline) <= 45
            and len(cleaned_inline.split()) <= 4
            and not re.search(r"\d", cleaned_inline)
        ):
            append_current_block()
            current_school = normalize_school_name(cleaned_inline)
            current_grade = None
            current_teachers = None
            current_feedback = []
            continue

        start_match = start_date_pattern.search(line)
        end_match = end_date_pattern.search(line)
        date_match = date_pattern.search(line)
        if start_match or end_match or date_match:
            append_current_block()
            current_grade = None
            current_teachers = None
            current_feedback = []
            current_notes_count = None
            current_reported_by = None

            if start_match:
                dt = parse_date(start_match.group(1))
                if dt:
                    current_start_date = dt.strftime("%Y-%m-%d")
            if end_match:
                dt = parse_date(end_match.group(1))
                if dt:
                    days_to_subtract = (dt.weekday() + 1) % 7
                    sunday_date = dt - timedelta(days=days_to_subtract)
                    current_end_date = sunday_date.strftime("%Y-%m-%d")
            if date_match and not (start_match or end_match):
                dt = parse_date(date_match.group(1))
                if dt:
                    days_to_subtract = (dt.weekday() + 1) % 7
                    sunday_date = dt - timedelta(days=days_to_subtract)
                    current_end_date = sunday_date.strftime("%Y-%m-%d")
                    current_start_date = (sunday_date - timedelta(days=6)).strftime("%Y-%m-%d")
            continue

        grade_match = grade_pattern.match(line)
        if grade_match:
            if len(line) > 120:
                if (current_start_date or current_end_date) and current_grade:
                    current_feedback.append(line)
                continue

            append_current_block()
            current_feedback = []
            level_type = (grade_match.group(1) or grade_match.group(4)).lower()
            grade_num = int(grade_match.group(2) or grade_match.group(3))
            current_grade = format_grade_or_year(grade_num, level_type)
            trailing_text = grade_match.group(5).strip()
            if trailing_text:
                current_teachers = trailing_text
            continue

        # relaxed grade parser when the strict header regex misses
        extracted_grade = extract_grade_from_text(line)
        if extracted_grade:
            append_current_block()
            current_feedback = []
            current_grade = extracted_grade
            continue

        if (current_start_date or current_end_date):
            current_feedback.append(line)

    append_current_block()

    if not data:
        print(f"WARNING: No rows extracted from {file_path}")
        return pd.DataFrame()

    columns = ["Start_Date", "End_Date_Sunday", "School_Name", "Grade_Or_Year", "Teacher_Names", "Feedback_Text"]
    return pd.DataFrame(data, columns=columns)


def add_period_columns(df):
    if df.empty:
        return df
    out = df.copy()
    out["Start_Date_dt"] = pd.to_datetime(out["Start_Date"], errors="coerce")
    out["End_Date_dt"] = pd.to_datetime(out["End_Date_Sunday"], errors="coerce")
    out["Sort_Date"] = out["End_Date_dt"].fillna(out["Start_Date_dt"])
    out["Period_Year"] = out["Sort_Date"].dt.year
    out["Half"] = out["Sort_Date"].dt.month.apply(
        lambda month: "H1" if pd.notna(month) and int(month) <= 6 else ("H2" if pd.notna(month) else "Unknown")
    )
    out["Half_Year_Sheet"] = out.apply(
        lambda row: f"{int(row['Period_Year'])}_{row['Half']}" if pd.notna(row["Period_Year"]) else "Unknown_Period",
        axis=1,
    )
    return out


def fill_missing_dates_from_filename(df, docx_files):
    if df.empty:
        return df
    out = df.copy()
    for idx, row in out.iterrows():
        need_start = pd.isna(row["Start_Date"]) or str(row["Start_Date"]).strip() == ""
        need_end = pd.isna(row["End_Date_Sunday"]) or str(row["End_Date_Sunday"]).strip() == ""
        if not (need_start or need_end):
            continue
        source_file = docx_files[int(idx)] if int(idx) < len(docx_files) else ""
        start_dt, end_dt = extract_date_range_from_filename(str(source_file))
        if need_start and start_dt:
            out.at[idx, "Start_Date"] = start_dt.strftime("%Y-%m-%d")
        if need_end and end_dt:
            days_to_subtract = (end_dt.weekday() + 1) % 7
            sunday_date = end_dt - timedelta(days=days_to_subtract)
            out.at[idx, "End_Date_Sunday"] = sunday_date.strftime("%Y-%m-%d")
    return out


def convert_feedback_folder_to_excel(input_folder, output_excel):
    input_path = Path(input_folder)
    if not input_path.exists():
        raise FileNotFoundError(f"Input folder does not exist: {input_folder}")

    docx_files = sorted(input_path.glob("*.docx"))
    if not docx_files:
        raise FileNotFoundError(f"No .docx files found in: {input_folder}")

    all_frames = []
    file_tags = []
    for file_path in docx_files:
        try:
            frame = process_feedback_text(file_path)
            if not frame.empty:
                all_frames.append(frame)
                file_tags.extend([file_path.name] * len(frame))
                print(f"Parsed {file_path.name}: {len(frame)} rows")
            else:
                print(f"No matched rows in {file_path.name}")
        except Exception as exc:
            print(f"Error in {file_path.name}: {exc}")

    if not all_frames:
        raise RuntimeError("No feedback data extracted from any document.")

    merged = pd.concat(all_frames, ignore_index=True)
    merged = fill_missing_dates_from_filename(merged, file_tags)
    merged = merged[
        merged["Start_Date"].notna() | merged["End_Date_Sunday"].notna()
    ].copy()
    merged["Grade_Or_Year"] = merged["Grade_Or_Year"].fillna("").astype(str).str.strip()
    merged = merged[~((merged["Grade_Or_Year"] == "") & (merged["Feedback_Text"].fillna("").str.strip() == ""))].copy()
    merged = add_period_columns(merged)
    merged = merged[merged["Half_Year_Sheet"] != "Unknown_Period"].copy()
    merged = merged[merged["Feedback_Text"].fillna("").str.strip() != ""].copy()

    # Fill missing teacher names from nearby rows of same School + Grade
    merged["Teacher_Names"] = merged["Teacher_Names"].fillna("").astype(str).str.strip()
    merged["Teacher_Names"] = (
        merged.groupby(["School_Name", "Grade_Or_Year"], dropna=False)["Teacher_Names"]
        .transform(lambda s: s.replace("", pd.NA).ffill().bfill().fillna(""))
    )

    merged["Grade_Sort"] = merged["Grade_Or_Year"].apply(grade_sort_key)

    merged.sort_values(
        by=["Sort_Date", "School_Name", "Grade_Sort", "Grade_Or_Year"],
        ascending=[True, True, True, True],
        inplace=True,
    )

    output_columns = [
        "Start_Date",
        "End_Date_Sunday",
        "School_Name",
        "Grade_Or_Year",
        "Teacher_Names",
        "Feedback_Text",
    ]

    # Excel display format requested: dd-mm-yyyy
    merged["Start_Date"] = merged["Start_Date"].apply(format_date_for_excel)
    merged["End_Date_Sunday"] = merged["End_Date_Sunday"].apply(format_date_for_excel)

    with pd.ExcelWriter(output_excel, engine="openpyxl") as writer:
        merged[output_columns].to_excel(writer, sheet_name="All_Data", index=False)
        for sheet_name in sorted(merged["Half_Year_Sheet"].dropna().unique()):
            period_df = merged[merged["Half_Year_Sheet"] == sheet_name].copy()
            period_df.sort_values(
                by=["School_Name", "Grade_Sort", "Sort_Date", "Grade_Or_Year"],
                ascending=[True, True, True, True],
                inplace=True,
            )
            period_df[output_columns].to_excel(writer, sheet_name=str(sheet_name)[:31], index=False)

    print(f"\nSuccess: {len(merged)} rows written -> {output_excel}")
    return merged


if __name__ == "__main__":
    INPUT_FOLDER = "/home/gogul/Documents/TSAP/teachers_support_feedback_ai/Teachers_notes"
    OUTPUT_EXCEL = "/home/gogul/Documents/TSAP/teachers_support_feedback_ai/backend/rag/Cleaned_Feedback_HalfYear_Sorted.xlsx"
    convert_feedback_folder_to_excel(INPUT_FOLDER, OUTPUT_EXCEL)